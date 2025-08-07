from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH


def remove_lateral_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for child in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(child)

    tcBorders = OxmlElement('w:tcBorders')
    
    
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '8')     
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '000000')
    tcBorders.append(top)

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    tcBorders.append(bottom)

    for side in ('left', 'right'):
        edge = OxmlElement(f'w:{side}')
        edge.set(qn('w:val'), 'nil')  
        tcBorders.append(edge)

    tcPr.append(tcBorders)

def addTitle(Document,Titulo):
    sectionTitleTBL = Document.add_table(rows=1,cols=2)
    sectionTitleTBL.columns[0].width = Inches(4.5)
    sectionTitleTBL.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    sectionTitleTBL.rows[0].height = Pt(18)
    sectionTitleTBL.alignment = WD_TABLE_ALIGNMENT.CENTER
    sectionTitleCell = sectionTitleTBL.cell(0,0)
    sectionTitleCell.merge(sectionTitleTBL.cell(0,1))
    sectionTitleCell.paragraphs[0].text = Titulo
    remove_lateral_borders(sectionTitleCell)
    run_title = sectionTitleCell.paragraphs[0].runs[0]
    run_title.font.bold = True
    run_title.font.size = Pt(14)
    run_title.font.name = "Arial"

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)  # Cor de fundo (ex: 'D9D9D9')
    
    for e in tcPr.findall(qn('w:shd')):
        tcPr.remove(e)
        
    tcPr.append(shd)

def addLine(Document,info,answer,color=False,Aligment='Center'):
    sectionLineTBL = Document.add_table(rows=1,cols=2)
    sectionLineTBL.alignment = WD_TABLE_ALIGNMENT.CENTER
    sectionLineCellQ = sectionLineTBL.cell(0,0)
    sectionLineTBL.columns[0].width= Inches(2.3)
    sectionLineCellQ.paragraphs[0].text = info
    sectionLineCellQ.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if Aligment == 'Start':
            sectionLineCellQ.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = sectionLineCellQ.paragraphs[0].runs[0]
    run_title.font.size = Pt(12)
    run_title.font.name = "Arial"
    
    sectionLineCellI = sectionLineTBL.cell(0,1)
    sectionLineTBL.columns[1].width= Inches(5.2)
    sectionLineCellI.paragraphs[0].text =  answer
    sectionLineCellI.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if Aligment == 'Start':
            sectionLineCellI.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = sectionLineCellI.paragraphs[0].runs[0]
    run_title.font.size = Pt(12)
    run_title.font.name = "Arial"

    if color==True:
        set_cell_background(sectionLineCellI,'d9d9d9')
        set_cell_background(sectionLineCellQ,'d9d9d9')

def addSubTitle(Document,Subtitle, Color = False):
    sectionSubTitleTBL = Document.add_table(rows=1,cols=2)
    sectionSubTitleTBL.alignment = WD_TABLE_ALIGNMENT.CENTER
    sectionSubTitleTBL.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    sectionSubTitleTBL.rows[0].height = Pt(16)
    sectionSubTitleCell = sectionSubTitleTBL.cell(0,0)
    sectionSubTitleTBL.columns[0].width = Inches(2.3)
    sectionSubTitleTBL.columns[1].width = Inches(5.2)
    sectionSubTitleCell.merge(sectionSubTitleTBL.cell(0,1))
    sectionSubTitleCell.paragraphs[0].text = Subtitle
    run_title = sectionSubTitleCell.paragraphs[0].runs[0]
    run_title.font.bold = True
    run_title.font.size = Pt(12)
    run_title.font.name = "Arial"
    
    if Color == True:
        set_cell_background(sectionSubTitleCell,'d9d9d9')
        
def addPictogram(Document,PictoPath, Width, Height):
    pictogramTBL = Document.add_table(rows=1,cols=1)
    pictogramTBL.columns[0].width = Inches(7.5)
    pictoCell =pictogramTBL.cell(0,0)
    pictoP = pictoCell.paragraphs[0]
    pictoP.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pictoP.add_run().add_picture(PictoPath, width=Inches(Width), height=Inches(Height))
        
def mkSec1(Document,ProductName,Uses,ProviderInfo,Emergency):
    addTitle(Document,'1 - IDENTIFICAÇÃO')
    addLine(Document,'Identificação do produto:',ProductName,True)
    addLine(Document,'Usos recomendados:',Uses)
    addLine(Document,'Detalhes do fornecedor:',ProviderInfo,True)
    addLine(Document,'Número do telefone de emergência:',Emergency)

def mkSec2(Document,Classfication,ClassSystem,OtherDangerous,PictoPath,pictoWidth,pictoHeight,warningWord,warningPhrases,worryPhrases):
    addTitle(Document,'2 - IDENTIFICAÇÃO DE PERIGOS')
    addLine(Document,'Classificação da substância ou mistura:',Classfication,True)
    addLine(Document,'Sistema de classificação utilizado:',ClassSystem)
    addLine(Document,'Outros perigos que não resultam em uma classificação:',OtherDangerous,True)
    addSubTitle(Document,'Elementos de rotulagem do GHS, incluindo frases de precaução')
    addLine(Document,'Pictogramas:','',True)
    addPictogram(Document,PictoPath,pictoWidth,pictoHeight)
    addLine(Document,'Palavra de advertência:',warningWord,True)
    addLine(Document,'Frases de perigo:',warningPhrases)
    addLine(Document,'Frases de precaução:',worryPhrases,True)
    
def mkSec3(Document,subOrMix,chemID,synonym,cas,impure):
    addTitle(Document,'3 - COMPOSIÇÃO E INFORMAÇÕES SOBRE OS INGREDIENTES')
    addSubTitle(Document,subOrMix)
    addLine(Document,'Identidade química:',chemID,True)
    addLine(Document,'Sinônimo:',synonym)
    addLine(Document,'Número de registro CAS:',cas,True)
    addLine(Document,'Impurezas que contribuem para o perigo:',impure)
    
def mkSec4(Document,inhalation,skin,eyes,intake,after,doctor):
    addTitle(Document,'4 - MEDIDAS DE PRIMEIROS-SOCORROS')
    addLine(Document,'Inalação:',inhalation,True)
    addLine(Document,'Contato com a pele:',skin)
    addLine(Document,'Contato com os olhos:',eyes,True)
    addLine(Document,'Ingestão:',intake)
    addLine(Document,'Sintomas e efeitos mais importantes, agudos ou tardios:',after,True)
    addLine(Document,'Notas para o médico:',doctor)
    
def mkSec5(Document,extinction,EspDangerous,firefighters):
    addTitle(Document,'5 - MEDIDAS DE COMBATE A INCÊNDIO')
    addLine(Document,'Meios de extinção:',extinction,True)
    addLine(Document,'Perigos específicos da mistura ou substância:',EspDangerous)
    addLine(Document,'Medidas de proteção especiais para a equipe de combate a incêndio:',firefighters,True)

def mkSec6(Document,NonEmergencyPP,EmergencyPP,environment,containmentClean):
    addTitle(Document,'6 - MEDIDAS DE CONTROLE PARA DERRAMAMENTO OU VAZAMENTO')
    addSubTitle(Document,'Precauções pessoais')
    addLine(Document,'Para o pessoal que não faz parte dos serviços de emergência:',NonEmergencyPP,True)
    addLine(Document,'Para pessoal de serviço de emergência:',EmergencyPP)
    addLine(Document,'Precauções ao meio ambiente:',environment,True)
    addLine(Document,'Métodos e materiais para contenção e limpeza:',containmentClean) 

def mkSec7(Document,SafeHandling,hygiene,FireExplosion,adqCondition,adqPackage,idqPackage):
    addTitle(Document,'7- MANUSEIO E ARMAZENAMENTO')
    addSubTitle(Document,'Medidas técnicas apropriadas para o manuseio')
    addLine(Document,'Precauções para manuseio seguro:',SafeHandling,True)
    addLine(Document,'Medidas de higiene:',hygiene)
    addSubTitle(Document,'Condições de armazenamento seguro, incluindo qualquer incompatibilidade')
    addLine(Document,'Prevenção de incêndio e explosão:',FireExplosion,True)
    addLine(Document,'Condições adequadas:',adqCondition)
    addLine(Document,'Materiais adequados para embalagem:',adqPackage,True)
    addLine(Document,'Materiais inadequados para embalagem:',idqPackage)

def mkSec8(Document,exposure,biology,otherLimits,engineeringCtrl,eyesFace,skinBody,Breathing,Termic):
    addTitle(Document,'8- CONTROLE DE EXPOSIÇÃO E PROTEÇÃO INDIVIDUAL')
    addSubTitle(Document,'Parâmetros de controle')
    addLine(Document,'Limites de exposição ocupacional:',exposure,True)
    addLine(Document,'Indicadores biológicos:',biology)
    addLine(Document,'Outros limites e valores:',otherLimits,True)
    addLine(Document,'Medidas de controle de engenharia:', engineeringCtrl)
    addSubTitle(Document,'Medidas de proteção pessoal')
    addLine(Document,'Proteção dos olhos/face:',eyesFace,True)
    addLine(Document,'Proteção da pele e do corpo:',skinBody)
    addLine(Document,'Proteção respiratória:',Breathing,True)
    addLine(Document,'Perigos térmicos:',Termic)

def mkSec9(Document,physical_state,color,odor,melting_point,boiling_point,flammability,explosive_limit,flash_point,auto_ignition_temperature,
           decomposition_temperature,pH,kinematic_viscosity,water_solubility,partition_coefficient,vapor_pressure,relative_density,relative_vapor_density,particle_characteristics,OtherInfo):
    addTitle(Document,'9- PROPRIEDADES FÍSICAS E QUÍMICAS')
    addLine(Document, 'Estado físico:', physical_state, True,Aligment='Start')
    addLine(Document, 'Cor:', color,Aligment='Start')
    addLine(Document, 'Odor e limite de odor:', odor, True,Aligment='Start')
    addLine(Document, 'Ponto de fusão/ponto de congelamento:', melting_point,Aligment='Start')
    addLine(Document, 'Ponto de ebulição ou ponto de ebulição inicial e faixa de ebulição:', boiling_point,True,Aligment='Start')
    addLine(Document, 'Inflamabilidade (sólido; líquidos e gás):', flammability,Aligment='Start')
    addLine(Document, 'Limite inferior/superior de inflamabilidade ou explosividade:', explosive_limit,True,Aligment='Start')
    addLine(Document, 'Ponto de fulgor:', flash_point,Aligment='Start')
    addLine(Document, 'Temperatura de autoignição:', auto_ignition_temperature,True,Aligment='Start')
    addLine(Document, 'Temperatura de decomposição:', decomposition_temperature,Aligment='Start')
    addLine(Document, 'pH:', pH,True,Aligment='Start')
    addLine(Document, 'Viscosidade cinemática:', kinematic_viscosity,Aligment='Start')
    addLine(Document, 'Solubilidade:', water_solubility,True,Aligment='Start')
    addLine(Document, 'Coeficiente de partição - noctanol/água:', partition_coefficient,Aligment='Start')
    addLine(Document, 'Pressão de vapor:', vapor_pressure,True,Aligment='Start')
    addLine(Document, 'Densidade relativa:', relative_density,Aligment='Start')
    addLine(Document, 'Densidade de vapor relativa:', relative_vapor_density,True,Aligment='Start')
    addLine(Document, 'Características das partículas (sólidos):', particle_characteristics,Aligment='Start')
    addLine(Document, 'Outras informações:', OtherInfo,True,Aligment='Start')

def mkSec10(Document,Stability,reactivity,possibility_of_dangerous_reactions,
conditions_to_avoid,incompatible_materials,hazardous_decomposition_products):
    addTitle(Document,'10 - ESTABILIDADE E REATIVIDADE')
    addLine(Document,'Estabilidade:',Stability,True)
    addLine(Document,'Reatividade:',reactivity)
    addLine(Document,'Possibilidade de reações perigosas:',possibility_of_dangerous_reactions,True)
    addLine(Document,'Condições a serem evitadas:',conditions_to_avoid)
    addLine(Document,'Materiais incompatíveis:',incompatible_materials,True)
    addLine(Document,'Produtos perigosos da decomposição:',hazardous_decomposition_products)



def mksec11(Document,acute_toxicity,skin_corrosion_irritation,serious_eye_damage_eye_irritation,
respiratory_skin_sensitization,germ_cell_mutagenicity,carcinogenicity,reproductive_toxicity,
stot_single_exposure,stot_repeated_exposure,aspiration_hazard):
    addTitle(Document,'11 - INFORMAÇÕES TOXICOLÓGICAS')
    addLine(Document,'Toxicidade aguda:',acute_toxicity,True)
    addLine(Document,'Corrosão/irritação à pele:',skin_corrosion_irritation)
    addLine(Document,'Lesões oculares graves/irritação ocular:',serious_eye_damage_eye_irritation,True)
    addLine(Document,'Sensibilização respiratória ou à pele:',respiratory_skin_sensitization)
    addLine(Document,'Mutagenicidade em células germinativas:',germ_cell_mutagenicity,True)
    addLine(Document,'Carcinogenicidade:',carcinogenicity)
    addLine(Document,'Toxicidade à reprodução:',reproductive_toxicity,True)
    addLine(Document,'Toxicidade para órgãos-alvo específicos - exposição única:',stot_single_exposure)
    addLine(Document,'Toxicidade para órgãos-alvo específicos - exposição repetida:',stot_repeated_exposure,True)
    addLine(Document,'Perigo por aspiração:',aspiration_hazard)


def mkSec12(Document,ecotoxicity,persistence_degradability,bioaccumulation_potential,
soil_mobility,other_adverse_effects):
    addTitle(Document,'12 - INFORMAÇÕES ECOLÓGICAS')
    addSubTitle(Document,'Efeitos ambientais, comportamento e impactos do produto')
    addLine(Document,'Ecotoxicidade:',ecotoxicity,True)
    addLine(Document,'Persistência e degradabilidade:',persistence_degradability)
    addLine(Document,'Potencial bioacumulativo:',bioaccumulation_potential,True)
    addLine(Document,'Mobilidade no solo:',soil_mobility)
    addLine(Document,'Outros efeitos adversos:',other_adverse_effects,True)


def mkSec13(Document,product_treatment,product_leftovers,used_packaging):
    addTitle(Document,'13- CONSIDERAÇÕES SOBRE DESTINAÇÃO FINAL')
    addSubTitle(Document,'Métodos recomendados para destinação final')
    addLine(Document,'Produto:',product_treatment,True)
    addLine(Document,'Restos de produtos:',product_leftovers)
    addLine(Document,'Embalagem usada:',used_packaging,True)

def mkSec14(Document,terrestrial,
            terrestrial_ONU,
            terrestrial_shipping_name,
            terrestrial_primary_class,
            terrestrial_subsidiary_class,
            terrestrial_risk_number,
            terrestrial_packing_group,
            hydroviario,
            hydroviario_ONU,
            hydroviario_shipping_name,
            hydroviario_primary_class,
            hydroviario_subsidiary_class,
            hydroviario_packing_group,
            hydroviario_ems,
            hydroviario_marine_pollutant,
            aereo,
            aereo_ONU,
            aereo_shipping_name,
            aereo_primary_class,
            aereo_subsidiary_class,
            aereo_packing_group):
    addTitle(Document,'14 - INFORMAÇÕES SOBRE TRANSPORTE')
    addSubTitle(Document,'Regulamentações nacionais e internacionais')
    addLine(Document,'Terrestre:',terrestrial,True)
    addLine(Document,'Número ONU:',terrestrial_ONU)
    addLine(Document,'Nome apropriado para embarque:',terrestrial_shipping_name,True)
    addLine(Document,'Classe ou subclasse de risco principal:',terrestrial_primary_class)
    addLine(Document,'Classe ou subclasse de risco subsidiário:',terrestrial_subsidiary_class,True)
    addLine(Document,'Número de risco:',terrestrial_risk_number)
    addLine(Document,'Grupo de embalagem:',terrestrial_packing_group,True)

    addLine(Document,'Hidroviário:', hydroviario)
    addLine(Document,'Número ONU:', hydroviario_ONU,True)
    addLine(Document,'Nome apropriado para embarque:', hydroviario_shipping_name)
    addLine(Document,'Classe ou subclasse de risco principal:', hydroviario_primary_class,True)
    addLine(Document,'Classe ou subclasse de risco subsidiário:', hydroviario_subsidiary_class)
    addLine(Document,'Grupo de embalagem:', hydroviario_packing_group,True)
    addLine(Document,'EmS:', hydroviario_ems)
    addLine(Document,'Poluente marinho:', hydroviario_marine_pollutant,True)

    addLine(Document,'Aéreo:', aereo)
    addLine(Document,'Número ONU:', aereo_ONU,True)
    addLine(Document,'Nome apropriado para embarque:', aereo_shipping_name)
    addLine(Document,'Classe ou subclasse de risco principal:', aereo_primary_class,True)
    addLine(Document,'Classe ou subclasse de risco subsidiário:', aereo_subsidiary_class)
    addLine(Document,'Grupo de embalagem:', aereo_packing_group,True)

def mkSec15(Document):
    legislacao = '''Decreto Federal nº 10.088, de 5 de novembro de 2019;
Norma ABNT-NBR 14725;
Norma Regulamentadora nº 26 (Sinalização de segurança), do Ministério do Trabalho e Emprego.'''
    addTitle(Document,'15 - INFORMAÇÕES SOBRE REGULAMENTAÇÕES')
    addLine(Document,'Regulamentações específicas para o produto químico:',legislacao,True)





























































