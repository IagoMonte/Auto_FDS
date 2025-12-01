from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE,WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

DEFAULT_FONT = "Arial"
DEFAULT_SIZE = Pt(12)

def formatRun(run, bold=False, size=DEFAULT_SIZE, name=DEFAULT_FONT):
    run.font.name = name
    run.font.size = size
    run.bold = bold

def newTable(doc, rows, cols, width_in=None, align=WD_TABLE_ALIGNMENT.CENTER):
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = align
    if width_in is not None:
        table.columns[0].width = Inches(width_in)
    return table

def setCellBackground(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)  #Cor de fundo (ex:'D9D9D9')
    
    for e in tcPr.findall(qn('w:shd')):
        tcPr.remove(e)
        
    tcPr.append(shd)

def removeLateralBorders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for child in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(child)

    tcBorders = OxmlElement('w:tcBorders')
    
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '8')     
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '000000')
    tcBorders.append(top)

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    tcBorders.append(bottom)

    for side in ('left', 'right'):
        edge = OxmlElement(f'w:{side}')
        edge.set(qn('w:val'), 'nil')  
        tcBorders.append(edge)

    tcPr.append(tcBorders)

def styleParagraph(p, bold=False, size=DEFAULT_SIZE, name=DEFAULT_FONT, align=WD_ALIGN_PARAGRAPH.CENTER):
    if p.runs:
        r = p.runs[0]
    else:
        r = p.add_run()
    formatRun(r, bold=bold, size=size, name=name)
    p.alignment = align


def addTitle(document, titulo):
    table = newTable(document, 1, 2, width_in=7.5)
    row = table.rows[0]
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    row.height = Pt(18)

    cell = table.cell(0, 0)
    cell.merge(table.cell(0, 1))
    p = cell.paragraphs[0]
    p.text = titulo
    styleParagraph(p, bold=True, size=Pt(14))
    removeLateralBorders(cell)

def addSubTitle(document, subtitle, color=False):
    table = newTable(document, 1, 2, width_in=7.5)
    row = table.rows[0]
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    row.height = Pt(16)

    table.columns[0].width = Inches(2.3)
    table.columns[1].width = Inches(5.2)

    cell = table.cell(0, 0)
    cell.merge(table.cell(0, 1))
    p = cell.paragraphs[0]
    p.text = subtitle
    styleParagraph(p, bold=True)

    if color:
        setCellBackground(cell, 'd9d9d9')

def addLine(document, info, answer, color=False, alignment='Center'):
    table = newTable(document, 1, 2, width_in=7.5)
    table.columns[0].width = Inches(2.3)
    table.columns[1].width = Inches(5.2)

    alignMap = {
        'Center': WD_ALIGN_PARAGRAPH.CENTER,
        'Start': WD_ALIGN_PARAGRAPH.LEFT
    }
    align = alignMap.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)

    # Coluna da pergunta
    cellQ = table.cell(0, 0)
    pQ = cellQ.paragraphs[0]
    pQ.text = info
    styleParagraph(pQ, align=align)

    # Coluna da resposta
    cellA = table.cell(0, 1)
    pA = cellA.paragraphs[0]
    pA.text = answer
    styleParagraph(pA, align=align)

    if color:
        setCellBackground(cellQ, 'd9d9d9')
        setCellBackground(cellA, 'd9d9d9')

    if answer == '':
        cellQ.merge(cellA)

def addVersionControl(document, version, date, info):
    outer = newTable(document, 1, 1, width_in=7.5)
    row = outer.rows[0]
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    row.height = Pt(132)
    cell = outer.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    tbl = cell.add_table(rows=2, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col in tbl.columns:
        col.width = Inches(7.0 / 3)  #~2.33

    for r_idx, h in enumerate((18, 45)):
        tbl.rows[r_idx].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        tbl.rows[r_idx].height = Pt(h)

    headers = ["Versão", "Data de elaboração", "Alterações"]
    values = [version, date, info]

    for col_idx, text in enumerate(headers):
        cell = tbl.cell(0, col_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.text = text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        formatRun(p.runs[0], bold=False)

    for col_idx, text in enumerate(values):
        cell = tbl.cell(1, col_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.text = text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        formatRun(p.runs[0], bold=False)

def addPictogram(Document,PictoPath, Width, Height):
    pictogramTBL = Document.add_table(rows=1,cols=1)
    pictogramTBL.columns[0].width = Inches(7.5)
    pictoCell =pictogramTBL.cell(0,0)
    pictoP = pictoCell.paragraphs[0]
    pictoP.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if PictoPath != None:
        pictoP.add_run().add_picture(PictoPath, width=Inches(Width), height=Inches(Height))
    else:
        pictoP.text = "Não disponivel"
        runTitle = pictoP.runs[0]
        formatRun(runTitle)