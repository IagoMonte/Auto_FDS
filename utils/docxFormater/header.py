from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def createFieldRun(run, field_code):
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = field_code
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    textRun = OxmlElement('w:t'); textRun.text = "1"
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar1, instrText, fldChar2, textRun, fldChar3])

def addPageNumber(paragraph):
    p = paragraph.add_run("Página "); p.bold = True; p.font.name = p.font.name = "Arial"; p.font.size = Pt(12)
    run = paragraph.add_run(); run.bold = True; run.font.name = "Arial"; run.font.size = Pt(12); createFieldRun(run, "PAGE")
    paragraph.add_run(" de ").bold = True; paragraph.runs[-1].font.name = "Arial"; paragraph.runs[-1].font.size = Pt(12)
    run2 = paragraph.add_run(); run2.bold = True; run2.font.name = "Arial"; run2.font.size = Pt(12); createFieldRun(run2, "NUMPAGES")

def HeaderGen(document, title):
    header = document.sections[0].header
    table = header.add_table(rows=4, cols=5, width=Pt(90))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    widths = [Inches(2.2), Inches(0.1), Inches(3), Inches(0.1), Inches(2.1)]
    for col, width in enumerate(widths): table.columns[col].width = width
    
    logoCell = table.cell(0, 0).merge(table.cell(3, 0))
    titleCell = table.cell(0, 2).merge(table.cell(3, 2))
    for col in [1, 3]: table.cell(0, col).merge(table.cell(3, col))
    
    logoCell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pLogo = logoCell.paragraphs[0]; pLogo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pLogo.add_run().add_picture('./src/LOGO.png', width=Inches(1.8))
    

    titleCell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pTitle = titleCell.paragraphs[0]; pTitle.text = f"FDS - {title}\t"; pTitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pTitle.runs[0].font.update(bold=True, size=Pt(18), name="Arial")
    

    rCells = [
        ('FOR-40', table.cell(0, 4)),
        ('Revisao: 01', table.cell(1, 4)),
        ('DATA:\n28/07/2024', table.cell(2, 4))
    ]
    
    for text, cell in rCells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.add_paragraph(text); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.update(bold=True, size=Pt(12), name="Arial")
    
    page_cell = table.cell(3, 4); page_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    addPageNumber(page_cell.add_paragraph())
    
    header.add_paragraph('')
    return document