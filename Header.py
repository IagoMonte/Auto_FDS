from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Cria elemento de bordas da célula
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'nil')  # sem borda
        tcBorders.append(edge_element)

    # Remove qualquer borda anterior e define sem borda
    for child in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(child)
    tcPr.append(tcBorders)

def add_page_number(paragraph):
    """Insere 'Página X de Y' com estilo Arial Bold."""

    # Texto inicial: "Página "
    run_page = paragraph.add_run("Página ")
    run_page.bold = True
    run_page.font.name = "Arial"
    run_page.font.size = Pt(12)

    # === Campo: PAGE ===
    run = paragraph.add_run()
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    text_run = OxmlElement('w:t')
    text_run.text = "1"  # valor temporário

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(text_run)
    run._r.append(fldChar3)

    # Texto intermediário: " de "
    run_de = paragraph.add_run(" de ")
    run_de.bold = True
    run_de.font.name = "Arial"
    run_de.font.size = Pt(12)

    # === Campo: NUMPAGES ===
    run2 = paragraph.add_run()
    run2.bold = True
    run2.font.name = "Arial"
    run2.font.size = Pt(12)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    text_run = OxmlElement('w:t')
    text_run.text = "1"

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    run2._r.append(text_run)
    run2._r.append(fldChar3)
    
    
def HeaderGen(document, title):
    doc = document
    header = doc.sections[0].header

    Headertbl = header.add_table(rows=4, cols=5, width=Pt(90))
    Headertbl.style = 'Table Grid'
    Headertbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    Headertbl.columns[0].width = Inches(2.3)
    Headertbl.columns[1].width = Inches(0.1)
    Headertbl.columns[2].width = Inches(3)
    Headertbl.columns[3].width = Inches(0.1)
    Headertbl.columns[4].width = Inches(2.3)  

    Headertbl.cell(0,1).merge(Headertbl.cell(1,1))# space1
    Headertbl.cell(0,1).merge(Headertbl.cell(2,1))# space1
    Headertbl.cell(0,1).merge(Headertbl.cell(3,1))# space1


    Headertbl.cell(0,2).merge(Headertbl.cell(1,2))# Title
    Headertbl.cell(0,2).merge(Headertbl.cell(2,2))# Title
    Headertbl.cell(0,2).merge(Headertbl.cell(3,2))# Title

    Headertbl.cell(0,3).merge(Headertbl.cell(1,3))# space2
    Headertbl.cell(0,3).merge(Headertbl.cell(2,3))# space2
    Headertbl.cell(0,3).merge(Headertbl.cell(3,3))# space2

    cell_logo = Headertbl.cell(0, 0)
    cell_logo.merge(Headertbl.cell(1,0))
    cell_logo.merge(Headertbl.cell(2,0))
    cell_logo.merge(Headertbl.cell(3,0))


    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run()
    run_logo.add_picture('./src/LOGO.png',width=Inches(1.8))

    cell_title = Headertbl.cell(0, 2)
    cell_title.merge(Headertbl.cell(1,2))
    cell_title.merge(Headertbl.cell(2,2))
    cell_title.merge(Headertbl.cell(3,2))


    cell_title.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_title = cell_title.paragraphs[0]
    p_title.text = f"FDS - {title}\t"
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.runs[0]
    run_title.font.bold = True
    run_title.font.size = Pt(18)
    run_title.font.name = "Arial"

    For_40Cell = Headertbl.cell(0,4)
    For_40Cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    RevCell = Headertbl.cell(1,4)
    RevCell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    DataCell = Headertbl.cell(2,4)
    DataCell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    PageCell = Headertbl.cell(3,4)
    PageCell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    For_40Celltxt = For_40Cell.add_paragraph('FOR-40')
    For_40Celltxt.runs[0].font.bold = True
    For_40Celltxt.runs[0].font.size = Pt(12)
    For_40Celltxt.runs[0].font.name = "Arial"
    For_40Celltxt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    RevCelltxt = RevCell.add_paragraph('Revisao: 01')
    RevCelltxt.runs[0].font.bold = True
    RevCelltxt.runs[0].font.size = Pt(12)
    RevCelltxt.runs[0].font.name = "Arial"
    RevCelltxt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    DataCelltxt = DataCell.add_paragraph('''DATA:\n28/07/2024''')
    DataCelltxt.runs[0].font.bold = True
    DataCelltxt.runs[0].font.size = Pt(12)
    DataCelltxt.runs[0].font.name = "Arial"
    DataCelltxt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    PageCelltxt = PageCell.add_paragraph()
    add_page_number(PageCelltxt)
    PageCelltxt.runs[0].font.bold = True
    PageCelltxt.runs[0].font.size = Pt(12)
    PageCelltxt.runs[0].font.name = "Arial"
    PageCelltxt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_paragraph('')
    
    return doc